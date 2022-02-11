import os
import psutil
from PyQt5.QtCore import Qt, QDate
from PyQt5.QtGui import QFont, QIcon
from PyQt5.QtWidgets import (QCheckBox, QGroupBox, QLabel, QLineEdit, QPushButton, QRadioButton,
                             QTabWidget, QVBoxLayout, QWidget, QMessageBox, QDateEdit, QFileDialog,
                             QScrollArea, QMainWindow, QTreeView, QHBoxLayout, QGridLayout)
from PyQt5.Qt import QStandardItemModel, QStandardItem
from mailmerge import MailMerge
from docx import Document
from docx.shared import Inches
from PyQt5 import uic
import json
import pandas as pd
from openpyxl import load_workbook


class CreateReportPage(QMainWindow):
    def __init__(self, parent):
        super(CreateReportPage, self).__init__()
        uic.loadUi('ui/createImplant.ui', self)
        self.parentWindow = parent
        # self.setGeometry(100, 60, 840, 670)
        self.showMaximized()

        self.setWindowTitle('Create Report')
        self.setWindowIcon(QIcon('data/favicon.ico'))

        self.setWindowFlags(
            Qt.Window |
            Qt.CustomizeWindowHint |
            Qt.WindowTitleHint |
            Qt.WindowCloseButtonHint
        )

        self.myFont = QFont("MS Shell Dlg 2", 12)
        self.myFontBold = QFont("MS Shell Dlg 2", 12, QFont.Bold)

        self.patientName = self.findChild(QLineEdit, 'patientName')
        self.chartNumber = self.findChild(QLineEdit, 'chartNumber')

        #for doctorIndex in [1,2,3]:
        #    self.findChild(QRadioButton, "doctor"+str(doctorIndex)).clicked.connect(self.doctorChanged)
        self.doctorBox = self.findChild(QVBoxLayout, "doctorBox")
        #print(self.doctorBox)

        #self.doctorBox.setExclusive(True)
        self.createDoctors()

        self.doctor = "David W. Engen, DDS, MSD"

        self.date = self.findChild(QDateEdit, 'date')
        self.uncoverDate = self.findChild(QDateEdit, 'uncoverDate')
        self.singleStage = self.findChild(QCheckBox, 'singleStage')
        self.restoreDate = self.findChild(QDateEdit, 'restoreDate')

        self.singleStage.clicked.connect(self.singleStagePressed)


        self.firstImplant = NewImplantForm()
        self.implantLayout = self.findChild(QVBoxLayout, "implantLayout")
        self.implantLayout.addWidget(self.firstImplant)

        self.implantCount = 1


        self.implantRestoreBox = self.findChild(QGroupBox, 'implantRestoreBox')

        # Disgusting code
        self.findChild(QLabel, 'restoreOpt3').mousePressEvent = self.restoreLabelConnectToBtn3
        self.findChild(QLabel, 'restoreOpt4').mousePressEvent = self.restoreLabelConnectToBtn4
        self.findChild(QLabel, 'restoreOpt5').mousePressEvent = self.restoreLabelConnectToBtn5

        self.anestheticLayout = self.findChild(QVBoxLayout, 'anestheticBox')
        self.anestheticBox = self.findChild(QGroupBox, 'anestheticGroupBox')
        self.toleranceBox = self.findChild(QGroupBox, 'toleranceBox')
        self.prescriptionsLayout = self.findChild(QVBoxLayout, 'prescriptionsBox')
        self.prescriptionsBox = self.findChild(QGroupBox, 'prescriptionsGroupBox')
        self.createAnesthetic()
        self.createRx()

        self.xrayButton = self.findChild(QPushButton, 'addXray')
        self.xrayButton.clicked.connect(self.getXrayPath)
        self.xrayPath = None

        self.addImplantButtonYes = self.findChild(QPushButton, 'addImplantYes')
        self.addImplantButtonYes.clicked.connect(self.addImplantTab)
        self.addImplantButtonNo = self.findChild(QPushButton, 'addImplantNo')
        # self.addImplantButtonNo.clicked.connect(self.something)

        self.tabWidget = self.findChild(QTabWidget, "createPages")

        # self.newImplantWidget = NewImplantForm()
        # print(self.newImplantWidget)
        # self.tabWidget.addTab(self.newImplantWidget, "Tab2")
        self.printSurgeryReportButton = self.findChild(QPushButton, "printSurgeryReport")
        self.printSurgeryReportButton.clicked.connect(self.printSurgeryReportPressed)


        #self.printChartNotesReportButton = self.findChild(QPushButton, "printChartNotesReport")
        #self.printChartNotesReportButton.clicked.connect(self.printChartNotesReportPressed)

        self.setDateDefaults()
        self.uncoverDate.dateChanged.connect(self.dateChanged)
        self.restoreDate.dateChanged.connect(self.dateChanged)
        self.dateChanged()
        self.tabWidget.setTabText(0,self.getTabName(self.implantCount))
        self.show()

    def singleStagePressed(self):
        if self.singleStage.isChecked():
            self.uncoverDate.setEnabled(False)
        else:
            self.uncoverDate.setEnabled(True)

    def setDateDefaults(self):
        self.date.setDate(QDate.currentDate())
        self.uncoverDate.setDate(QDate.currentDate().addMonths(6))
        self.restoreDate.setDate(QDate.currentDate().addMonths(6))

    def dateChanged(self):

        options = self.findChild(QGroupBox, "restoreTopBox").findChildren(QRadioButton)

        options[0].setText("Implant(s) can be restored after " + self.restoreDate.text() + ". ")
        options[1].setText("Implant(s) can be uncovered after " + self.uncoverDate.text() + " and restored after " + self.restoreDate.text() + ". ")


    def getXrayPath(self):
        with open("data/fileLocations.txt", "r") as content:
            lines = content.readlines()

        lastDir = " "
        if len(lines[1]) > 6:
            lastDir = lines[1][6:].strip()

        # print(lastDir)
        try:
            fname = QFileDialog.getOpenFileName(self, 'Open X-Ray', lastDir, "Image files (*.jpg *.jpeg)")
        except:
            fname = QFileDialog.getOpenFileName(self, 'Open X-Ray', "Image files (*.jpg *.jpeg)")
        if len(fname[0]) < 2:
            return
        self.xrayPath = fname[0]
        self.findChild(QLabel, 'xrayPath').setText("File Path: "+fname[0])

        with open("data/fileLocations.txt", "w") as content:
            content.write(lines[0])
            content.write("xrays="+os.path.dirname(fname[0])+"\n")
            content.write(lines[2])



    def addImplantTab(self):
        self.implantCount += 1
        newImplantWidget = NewImplantForm()
        self.tabWidget.addTab(newImplantWidget, self.getTabName(self.implantCount))

        self.tabWidget.setCurrentIndex(self.implantCount-1)
        self.findChild(QScrollArea, 'mainScroll').verticalScrollBar().setValue(0)

    def createDoctors(self):

        with open("data/doctors.txt", "r") as content:
            lines = content.readlines()
        docs = []
        for line in lines:
            if len(line) > 3:
                docs.append(line)

        self.docCount = len(docs)

        docGroup = QGroupBox()
        docLayout = QVBoxLayout()
        i = 0
        for doc in docs:
            d = QRadioButton(doc.replace("\n", ""))
            if i == 0:
                d.setChecked(True)
            d.setObjectName("doctor" + str(i))
            d.clicked.connect(self.doctorChanged)
            docLayout.addWidget(d)
            i += 1
        docGroup.setLayout(docLayout)
        self.doctorBox.addWidget(docGroup)

    def doctorChanged(self):
        for doctorIndex in range(self.docCount):
            doc = self.findChild(QRadioButton, "doctor" + str(doctorIndex))
            if doc.isChecked():
                self.doctor = doc.text()

        for i in range(self.tabWidget.count()):

            self.tabWidget.setTabText(i,self.getTabName(i+1))

    def createAnesthetic(self):
        bundle_dir = os.path.dirname(os.path.abspath(os.path.dirname(__file__)))
        path_to_dat = os.path.join(bundle_dir, 'data\\anesthetics.txt')
        with open(path_to_dat, "r") as content:
            lines = content.readlines()

        anesthetics = []
        for line in lines:
            if len(line) > 3:
                anesthetics.append(line.strip())
        # print(anesthetics)

        for anes in anesthetics:
            btn = QRadioButton(anes, font=self.myFont)
            self.anestheticLayout.addWidget(btn)

    def createRx(self):
        bundle_dir = os.path.dirname(os.path.abspath(os.path.dirname(__file__)))
        path_to_dat = os.path.join(bundle_dir, 'data\\prescriptions.txt')
        with open(path_to_dat, "r") as content:
            lines = content.readlines()

        prescriptions = []
        for line in lines:
            if len(line) > 3:
                prescriptions.append(line.strip())
        # print(prescriptions)

        for rx in prescriptions:
            box = QCheckBox(rx, font=self.myFont)
            self.prescriptionsLayout.addWidget(box)

    def getTabName(self, pageNum):
        return (self.doctor.split()[0]+"_Implant_"+str(pageNum))

    # Disgusting Code
    def restoreLabelConnectToBtn3(self, event):
        self.findChild(QRadioButton, 'restoreBtn3').setChecked(True)
        event.accept()
    def restoreLabelConnectToBtn4(self, event):
        self.findChild(QRadioButton, 'restoreBtn4').setChecked(True)
        event.accept()
    def restoreLabelConnectToBtn5(self, event):
        self.findChild(QRadioButton, 'restoreBtn5').setChecked(True)
        event.accept()

    def getRestore(self):
        if not self.implantRestoreBox.isChecked():
            return ""
        txt = ""

        topBox = self.findChild(QGroupBox, "restoreTopBox")
        botBox = self.findChild(QGroupBox, 'restoreBotBox')
        for opt in topBox.findChildren(QRadioButton):
            if opt.isChecked():
                txt += opt.text()
        for opt in botBox.findChildren(QRadioButton):
            if opt.isChecked():
                val = opt.objectName()[-1]
                txt += self.implantRestoreBox.findChild(QLabel, ("restoreOpt"+str(val))).text()

        return txt

    def getAnesthetic(self):
        if not self.anestheticBox.isChecked():
            return "N/A"
        options = self.anestheticBox.findChildren(QRadioButton)
        for opt in options:
            if opt.isChecked():
                return opt.text()
        return "N/A"

    def getTolerance(self):
        options = self.toleranceBox.findChildren(QRadioButton)
        for opt in options:
            if opt.isChecked():
                return opt.text()
        return "N/A"

    def getRX(self):
        if not self.prescriptionsBox.isChecked():
            return "N/A"
        options = self.prescriptionsBox.findChildren(QCheckBox)
        txt = ""
        for opt in options:
            if opt.isChecked():
                txt += (opt.text()+"\n")
        if txt == "":
            return "N/A"
        else:
            return txt

    def printSurgeryReportPressed(self):

        # print(self.firstImplant.generateParagraph())
        # print(self.firstImplant.getRestorativeParts())
        implants = ""
        reports = ""
        parts = ""
        healingCaps = ""
        for tab in self.tabWidget.findChildren(NewImplantForm):
            implants += (tab.getImplant()+"\n")
            reports += (tab.generateParagraph()+"\n\n")
            partsString = (tab.getRestorativeParts())
            if partsString not in parts:
                parts += partsString
            healingCaps += (tab.getHealingCapList()+"\n")

        restoreChoice = self.getRestore()
        anestheticChoice = self.getAnesthetic()
        toleranceChoice = self.getTolerance()
        rxChoice = self.getRX()
        if self.singleStage.isChecked():
            uncoverVal = "Single Stage"
        else:
            uncoverVal = self.date.text()
        template = "data/template.docx"

        document = MailMerge(template)
        # print(document.get_merge_fields())
        document.merge(
            patient=self.patientName.text(),
            chart=self.chartNumber.text(),
            surgeon=self.doctor,
            date=self.date.text(),
            uncover_date=uncoverVal,
            restore_date=self.restoreDate.text(),
            implant=implants,
            healing_cap=healingCaps,
            restorative_parts=parts,
            report=reports,
            restore=restoreChoice,
            anesthetic=anestheticChoice,
            tolerance=toleranceChoice,
            prescriptions=rxChoice)

        #xray = "C:\\PythonProjects\\ImplantApp\\app\\static\\images\\exampleXray.jpeg"
        document.write('temp.docx')
        doc = Document('temp.docx')
        if self.xrayPath != None:
            tables = doc.tables
            p = tables[0].rows[0].cells[0].add_paragraph()
            r = p.add_run()
            r.add_picture(self.xrayPath, width=Inches(2.5), height=Inches(2.5))

        firstName, lastName = self.separatePatient(self.patientName.text())
        dateString = self.date.dateTime().toString("yyyy_MM_dd")
        filename = lastName + "_" + firstName + "_" + dateString
        filename = filename.upper()
        filename = filename + ".docx"

        with open("data/fileLocations.txt", "r") as content:
            lines = content.readlines()
        dir = lines[0][8:].strip()
        if len(dir) < 3:
            dir = self.setDefaultFolder()
        pathName = os.path.join(dir, filename)
        if not os.path.exists(dir):
            dir = self.setDefaultFolder()
            pathName = os.path.join(dir, filename)
        doc.save(pathName)


        os.remove('temp.docx')

        self.printChartNotesReportPressed()
        os.startfile(pathName)
        self.closeAndOpenHome()

    def printChartNotesReportPressed(self):
        if "EXCEL.EXE" in (p.name() for p in psutil.process_iter()):
            error_dialog = QMessageBox()
            error_dialog.setWindowTitle("Error")
            error_dialog.setText('The Excel File is open. New rows cannot be created. They will have to be inputted manually. ')
            error_dialog.exec_()
            return

        # Open Excel Data
        with open("data/fileLocations.txt", "r") as content:
            lines = content.readlines()
        excelFile = lines[2][6:].strip()
        if len(excelFile) < 3:
            error_dialog = QMessageBox()
            error_dialog.setWindowTitle("Select File")
            error_dialog.setText("Select an Excel File. ")
            error_dialog.exec_()
            self.setDefaultExcel()
            with open("data/fileLocations.txt", "r") as content:
                lines = content.readlines()
            excelFile = lines[2][6:].strip()
        try:
            df = pd.read_excel(excelFile)
            if df.size == 0:
                error_dialog = QMessageBox()
                error_dialog.setWindowTitle("Error")
                error_dialog.setText('Cannot read the Excel File. New rows cannot be created. They will have to be inputted manually. ')
                error_dialog.exec_()
                return
        except:
            error_dialog = QMessageBox()
            error_dialog.setWindowTitle("Error")
            error_dialog.setText('The Excel File could not be found. New rows cannot be created. They will have to be inputted manually. ')
            error_dialog.exec_()
            return

        # Names and filling/removing cells
        columnNames = ["Patient", "", "Chart", "ImplantDate", "UncoverDate", "ExposedDate", "Details",
                       "NumberPlaced", "MinisPlaced"]
        df.drop(df.columns[len(columnNames):], 1, inplace=True)
        df.columns = columnNames
        df.replace("", float("NaN"), inplace=True)
        df.dropna(subset=["ImplantDate"], inplace=True)
        df["Patient"].fillna(method='ffill', inplace=True)
        df.replace(float("NaN"), "", inplace=True)

        # Test: Fields, Gene
        firstName, lastName = self.separatePatient(self.patientName.text())
        formattedName = lastName +", "+ firstName
        patientRecords = df[df['Patient'].str.match(formattedName)]
        # print(patientRecords)

        # Determine the insert index
        if not patientRecords.empty:
            lastIndex = int(patientRecords.index[-1])+3
        else:
            lastIndex = int(df.iloc[-1].name)+4
        # print(lastIndex)

        wb = load_workbook(excelFile)
        ws = wb.active
        rowCount = ws.max_row


        newRows = self.tabWidget.count()
        ws.insert_rows(lastIndex, newRows)

        currentIndex = lastIndex
        #print(type(ws.cell(row=6304, column=5).data_type))
        if patientRecords.empty:
            ws.cell(row=currentIndex, column=1).value = formattedName
            ws.cell(row=currentIndex, column=3).value = self.chartNumber.text()


        for tab in self.tabWidget.findChildren(NewImplantForm):
            # ws.cell(row=currentIndex, column=2).value = 'x'
            ws.cell(row=currentIndex, column=4).value = self.date.text()
            ws.cell(row=currentIndex, column=4).number_format = 'mm-dd-yy'
            if self.singleStage.isChecked():
                uncoverVal = "Single Stage"
            else:
                uncoverVal = self.date.text()
            ws.cell(row=currentIndex, column=5).value = uncoverVal
            ws.cell(row=currentIndex, column=5).number_format = 'mm-dd-yy'

            detailsString = tab.getDetails()
            ws.cell(row=currentIndex, column=7).value = detailsString
            ws.cell(row=currentIndex, column=8).value = "1"
            ws.cell(row=currentIndex, column=8).data_type = 'int64'
            currentIndex += 1

        print(lastIndex, rowCount)
        ws.delete_rows((rowCount-3-newRows), newRows)

        wb.save(excelFile)


    def separatePatient(self, name):

        firstName = ""
        lastName = ""
        if "," in name:
            parts = name.split(",", 1)
            lastName = parts[0].strip()
            firstName = parts[1].strip()

        elif " " in name:
            parts = name.split(" ", 1)
            firstName = parts[0].strip()
            lastName = parts[1].strip()
        else:
            lastName = name

        return firstName, lastName

    def setDefaultFolder(self):
        error_dialog = QMessageBox()
        error_dialog.setWindowTitle("Select Report Folder")
        error_dialog.setText("Please select a report folder. ")
        error_dialog.exec_()

        bundle_dir = os.path.dirname(os.path.abspath(os.path.dirname(__file__)))
        path_to_dat = os.path.join(bundle_dir, 'data\\fileLocations.txt')
        with open(path_to_dat, "r") as content:
            lines = content.readlines()

        file = str(QFileDialog.getExistingDirectory(self, "Select Directory"))
        if len(file) < 2:
            return "C:/"

        with open("data/fileLocations.txt", "w") as content:
            content.write("reports="+file+"\n")
            content.write(lines[1])
            content.write(lines[2])

        return file

    def setDefaultExcel(self):

        bundle_dir = os.path.dirname(os.path.abspath(os.path.dirname(__file__)))
        path_to_dat = os.path.join(bundle_dir, 'data\\fileLocations.txt')
        with open(path_to_dat, "r") as content:
            lines = content.readlines()

        path = lines[2][6:]
        if len(path) > 3:
            try:
                lastDir = os.path.dirname(path)
                file = QFileDialog.getOpenFileName(self, 'Open Excel', lastDir, "Microsoft Excel files (*.xlsx)")
            except:
                file = QFileDialog.getOpenFileName(self, 'Open Excel', filter="Microsoft Excel files (*.xlsx)")
        else:
            file = QFileDialog.getOpenFileName(self, 'Open Excel', filter="Microsoft Excel files (*.xlsx)")


        if len(file[0]) < 2:
            return
        # print(file)

        with open(path_to_dat, "w") as content:
            content.write(lines[0])
            content.write(lines[1])
            content.write("excel=" + file[0] + "\n")

    def closeAndOpenHome(self):
        self.parentWindow.show()
        self.close()

    def closeEvent(self, event):
        self.parentWindow.show()
        event.accept()



class NewImplantForm(QWidget):
    def __init__(self):
        super(NewImplantForm, self).__init__()
        uic.loadUi('ui/newImplantForm.ui', self)

        self.myFont = QFont("MS Shell Dlg 2", 12)
        self.myFontBold = QFont("MS Shell Dlg 2", 12, QFont.Bold)

        self.implant = ""
        self.implantTree = self.findChild(QTreeView, 'implantList')
        self.makeImplantTree()
        self.implantTree.clicked.connect(self.implantChanged)
        self.isMiniImplant = False

        self.toothNumber = self.findChild(QLineEdit, 'toothNumber')
        self.lotNumber = self.findChild(QLineEdit, 'lotNumber')
        self.expirationDate = self.findChild(QDateEdit, 'expirationDate')
        self.expirationDate.setDate(QDate.currentDate())

        self.restorativePartsList = self.findChild(QVBoxLayout, 'restorativePartsList')
        self.fillRestorativeParts()

        self.incisionBox = self.findChild(QGroupBox, 'incisionBox')
        self.incisionEdit1 = self.findChild(QLineEdit, 'incisionEdit1')
        self.incisionEdit2 = self.findChild(QLineEdit, 'incisionEdit2')
        self.incisionEdit1.textEdited.connect(self.incisionChanging)
        self.incisionEdit2.textEdited.connect(self.incisionChanging)
        self.incisionEdit1.editingFinished.connect(self.incisionChanged)
        self.incisionEdit2.editingFinished.connect(self.incisionChanged)

        self.flapBox = self.findChild(QGroupBox, 'flapBox')
        self.extractionBox = self.findChild(QGroupBox, 'extractionBox')
        self.extractionEdit = self.findChild(QLineEdit, 'extractionEdit')
        self.extractionEdit.textEdited.connect(self.extractionChanging)
        self.extractionEdit.editingFinished.connect(self.extractionChanged)

        self.osteotomyBox = self.findChild(QGroupBox, 'osteotomyBox')
        self.osteotomyEdit1 = self.findChild(QLineEdit, 'osteotomyEdit1')
        self.osteotomyEdit2 = self.findChild(QLineEdit, 'osteotomyEdit2')
        self.osteotomyEdit1.textEdited.connect(self.osteotomyChanging)
        self.osteotomyEdit2.textEdited.connect(self.osteotomyChanging)
        self.osteotomyEdit1.editingFinished.connect(self.osteotomyChanged)
        self.osteotomyEdit2.editingFinished.connect(self.osteotomyChanged)

        self.boneDensityBox = self.findChild(QGroupBox, 'boneDensityBox')
        self.perforationsBox = self.findChild(QGroupBox, 'perforationsBox')
        self.tappedBox = self.findChild(QGroupBox, 'tappedBox')
        self.sinusBox = self.findChild(QGroupBox, 'sinusBox')
        self.implantSeatedBox = self.findChild(QGroupBox, 'implantSeatedBox')
        self.depthBox = self.findChild(QGroupBox, 'depthBox')

        self.healingCapBox = self.findChild(QGroupBox, 'healingCapBox')
        self.healingCapEdit = self.findChild(QLineEdit, 'healingCapEdit')

        self.graftBox = self.findChild(QGroupBox, 'graftBox')
        self.createGrafts()

        #self.generateParagraph()


    def makeImplantTree(self):

        with open("data/implants.json", "r") as content:
            implants = json.load(content)

        treeModel = QStandardItemModel()
        rootNode = treeModel.invisibleRootItem()

        for root, child in implants.items():
            # print(root)
            outer = QStandardItem(root)
            # print(child)
            try:
                for root2, grandchild in child.items():
                    inner = QStandardItem(root2)

                    for gc in grandchild:
                        leaf = QStandardItem(gc)
                        inner.appendRow(leaf)
                    outer.appendRow(inner)

                    # print("\t"+str(root2))
                    # print("\t"+str(grandchild))
            except AttributeError:
                for c in child:
                    leaf = QStandardItem(c)
                    outer.appendRow(leaf)
            rootNode.appendRow(outer)

        self.implantTree.setModel(treeModel)
        self.implantTree.setHeaderHidden(True)

    def implantChanged(self, index):
        selectedItem = self.implantTree.model().itemFromIndex(index)
        self.implant = selectedItem.text()

        self.findChild(QLabel, "implantBoxLabel").setText("Implant: "+self.implant)

        firstButton = True
        for opt in self.implantSeatedBox.findChildren(QRadioButton):
            if firstButton:
                opt.setText("A "+ self.implant+" implant was seated to depth. ")
                firstButton = False
            else:
                opt.setText("A " + self.implant + " implant was not seated to depth. ")

    def fillRestorativeParts(self):

        def toggleRestorativeParts():
            partCount = self.restorativePartsList.count()
            cb = self.restorativePartsList.itemAt(0).widget()

            for i in range(1, partCount):
                if cb.isChecked():
                    self.restorativePartsList.itemAt(i).widget().setEnabled(False)
                else:
                    self.restorativePartsList.itemAt(i).widget().setEnabled(True)

        f = open("data/restorativeParts.txt", "r")
        myFont = QFont("MS Shell Dlg 2", 12)
        myFontBold = QFont("MS Shell Dlg 2", 12, QFont.Bold)
        noPartsOption = QCheckBox("No restorative parts to order. We will scan the implant digitally. "
                                  "You can choose from one of our partner Digital Implant Solutions labs "
                                  "who will make the models, custom abutment, and the restoration all "
                                  "according to your instructions. You can communicate with your lab directly.\n")
        noPartsOption.setChecked(True)
        noPartsOption.setFont(myFont)
        noPartsOption.toggled.connect(toggleRestorativeParts)
        self.restorativePartsList.addWidget(noPartsOption)
        # print(noPartsOption)

        for line in f:
            if line[0] == "&":


                category = QLabel()
                category.setText("\n" + line[1:-1])
                category.setFont(myFontBold)

                self.restorativePartsList.addWidget(category)

            else:
                if len(line) > 3:
                    option = QCheckBox()
                    option.setText(line[0:-1])
                    option.setFont(myFont)
                    self.restorativePartsList.addWidget(option)
        f.close()
        toggleRestorativeParts()

    def incisionChanging(self):

        if self.incisionBox.findChild(QLabel, 'helpMessage') == None:
            helpMessage = QLabel("Separate multiple teeth with commas.")
            helpMessage.setStyleSheet("color: red")
            helpMessage.setObjectName("helpMessage")

            self.incisionBox.layout().addWidget(helpMessage, 5, 3)

    def incisionChanged(self):
        helpMessage = self.incisionBox.findChild(QLabel, 'helpMessage')
        if helpMessage != None:
            self.incisionBox.layout().removeWidget(helpMessage)
            helpMessage.destroy()

    def extractionChanging(self):
        if self.extractionBox.findChild(QLabel, 'helpMessage') == None:
            helpMessage = QLabel("Separate multiple teeth with commas.")
            helpMessage.setStyleSheet("color: red")
            helpMessage.setObjectName("helpMessage")

            self.extractionBox.layout().addWidget(helpMessage, 2, 3)

    def extractionChanged(self):
        helpMessage = self.extractionBox.findChild(QLabel, 'helpMessage')
        if helpMessage != None:
            self.extractionBox.layout().removeWidget(helpMessage)
            helpMessage.destroy()

    def osteotomyChanging(self):
        if self.osteotomyBox.findChild(QLabel, 'helpMessage') == None:
            helpMessage = QLabel("Separate multiple teeth with commas.")
            helpMessage.setStyleSheet("color: red")
            helpMessage.setObjectName("helpMessage")

            self.osteotomyBox.layout().addWidget(helpMessage, 7, 0)

    def osteotomyChanged(self):
        helpMessage = self.osteotomyBox.findChild(QLabel, 'helpMessage')
        if helpMessage != None:
            self.osteotomyBox.layout().removeWidget(helpMessage)
            helpMessage.destroy()

    def createGrafts(self):
        bundle_dir = os.path.dirname(os.path.abspath(os.path.dirname(__file__)))
        path_to_dat = os.path.join(bundle_dir, 'data\\materials.txt')
        with open(path_to_dat, "r") as content:
            lines = content.readlines()

        self.materials = []
        self.membranes = []

        for line in lines:
            parts = line.split("_", 1)
            if len(parts) < 2:
                continue
            if parts[0] == "mat":
                self.materials.append(parts[1].strip())
            elif parts[0] == "mem":
                self.membranes.append(parts[1].strip())


        boneMaterialBox = self.graftBox.findChild(QVBoxLayout, "materialBox")
        boneMembraneBox = self.graftBox.findChild(QVBoxLayout, "membraneBox")

        for mat in self.materials:
            btn = QRadioButton(mat, font=self.myFont)
            boneMaterialBox.addWidget(btn)

        for mem in self.membranes:
            btn = QRadioButton(mem, font=self.myFont)
            boneMembraneBox.addWidget(btn)

    def getIncision(self):
        if not self.incisionBox.isChecked():
            return ""

        button1 = self.incisionBox.findChild(QRadioButton, 'incisionButton1')
        button2 = self.incisionBox.findChild(QRadioButton, 'incisionButton2')
        edit1 = self.incisionBox.findChild(QLineEdit, 'incisionEdit1')
        edit2 = self.incisionBox.findChild(QLineEdit, 'incisionEdit2')

        if button1.isChecked():
            multiple = len(edit1.text().split(",")) > 1 or len(edit1.text().split("-")) > 1
            if multiple:
                return "Sulcular incisions were made around teeth #"+edit1.text()+". "
            else:
                return "Sulcular incisions were made around tooth #" + edit1.text()+". "
        elif button2.isChecked():
            multiple = len(edit2.text().split(",")) > 1 or len(edit2.text().split("-")) > 1
            if multiple:
                return "A crestal incision was made distal of teeth #" + edit2.text()+". "
            else:
                return "A crestal incision was made between tooth #" + edit2.text()+". "
        else:
            return ""

    def getFlap(self):
        if not self.flapBox.isChecked():
            return ""

        for flap in self.flapBox.findChildren(QRadioButton):
            if flap.isChecked():
                return flap.text()
        return ""

    def getExtraction(self):
        if not self.extractionBox.isChecked():
            return ""

        edit = self.extractionBox.findChild(QLineEdit, "extractionEdit")
        btn1 = self.extractionBox.findChild(QRadioButton, "extractionButton1")
        btn2 = self.extractionBox.findChild(QRadioButton, "extractionButton2")
        btn3 = self.extractionBox.findChild(QRadioButton, "extractionButton3")
        btn4 = self.extractionBox.findChild(QRadioButton, "extractionButton4")


        multiple = len(edit.text().split(",")) > 1 or len(edit.text().split("-")) > 1
        txt = ""
        if multiple:
            if btn1.isChecked():
                txt += "Teeth " + edit.text() + " were removed via forcep extraction. "
                txt += "The sockets were perfectly preserved. "
            elif btn2.isChecked():
                txt += "Teeth " + edit.text() + " were removed via forcep extraction. "
                txt += "The sockets were well preserved. "
            elif btn3.isChecked():
                txt += "Teeth " + edit.text() + " were removed via forcep extraction. "
                txt += "The sockets were moderately preserved. "
            elif btn4.isChecked():
                txt += "Teeth " + edit.text() + " were removed via forcep extraction. "
                txt += "There was damage to the sockets. "
        else:
            if btn1.isChecked():
                txt += "Tooth " + edit.text() + " was removed via forcep extraction. "
                txt += "The socket was perfectly preserved. "
            elif btn2.isChecked():
                txt += "Tooth " + edit.text() + " was removed via forcep extraction. "
                txt += "The socket was well preserved. "
            elif btn3.isChecked():
                txt += "Tooth " + edit.text() + " was removed via forcep extraction. "
                txt += "The socket was moderately preserved. "
            elif btn4.isChecked():
                txt += "Tooth " + edit.text() + " was removed via forcep extraction. "
                txt += "There was damage to the socket. "
        return txt

    def getOsteotomy(self):
        if not self.osteotomyBox.isChecked():
            return ""

        button1 = self.osteotomyBox.findChild(QRadioButton, 'osteotomyButton1')
        button2 = self.osteotomyBox.findChild(QRadioButton, 'osteotomyButton2')
        edit1 = self.osteotomyBox.findChild(QLineEdit, 'osteotomyEdit1')
        edit2 = self.osteotomyBox.findChild(QLineEdit, 'osteotomyEdit2')

        if button1.isChecked():
            multiple = len(edit1.text().split(",")) > 1 or len(edit1.text().split("-")) > 1
            if multiple:
                return "An osteotomy was prepared in sites #"+edit1.text()+". "
            else:
                return "An osteotomy was prepared in site #" + edit1.text()+". "
        elif button2.isChecked():
            multiple = len(edit1.text().split(",")) > 1 or len(edit1.text().split("-")) > 1
            if multiple:
                return "An osteotomy was prepared in sockets #" + edit2.text()+". "
            else:
                return "An osteotomy was prepared in socket #" + edit2.text()+". "
        else:
            return ""

    def getBoneDensity(self):
        if not self.boneDensityBox.isChecked():
            return ""

        for density in self.boneDensityBox.findChildren(QRadioButton):
            if density.isChecked():
                return density.text()
        return ""

    def getPerforation(self):
        if not self.perforationsBox.isChecked():
            return ""

        for perforation in self.perforationsBox.findChildren(QRadioButton):
            if perforation.isChecked():
                return perforation.text()
        return ""

    def getTapped(self):
        if not self.tappedBox.isChecked():
            return ""

        for opt in self.tappedBox.findChildren(QRadioButton):
            if opt.isChecked():
                return opt.text()
        return ""

    def getSinus(self):
        if not self.sinusBox.isChecked():
            return ""

        for opt in self.sinusBox.findChildren(QRadioButton):
            if opt.isChecked():
                return opt.text()
        return ""

    def getSeated(self):
        if not self.implantSeatedBox.isChecked():
            return ""

        for opt in self.implantSeatedBox.findChildren(QRadioButton):
            if opt.isChecked():
                return opt.text()
        return ""

    def getDepth(self):
        if not self.depthBox.isChecked():
            return ""

        for opt in self.depthBox.findChildren(QRadioButton):
            if opt.isChecked():
                return opt.text()
        return ""

    def getHealingCap(self):
        if not self.healingCapBox.isChecked():
            return ""

        btn1 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton1')
        btn2 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton2')
        edit = self.healingCapBox.findChild(QLineEdit, 'healingCapEdit')


        if btn1.isChecked():
            return "A cover screw was placed. "
        elif btn2.isChecked():
            return "A " + edit.text() + " healing cap was placed for a single stage approach. "
        else:
            return ""

    def getGraft(self):
        if not self.graftBox.isChecked():
            return ""

        boneMaterialBox = self.findChild(QGroupBox, "boneMaterialBox")
        boneMembraneBox = self.findChild(QGroupBox, "boneMembraneBox")
        boneMaterial = ""
        boneMembrane = ""

        optionsChecked = False
        for opt in boneMaterialBox.findChildren(QRadioButton):
            if opt.isChecked():
                boneMaterial = opt.text()
                optionsChecked = True
        for opt in boneMembraneBox.findChildren(QRadioButton):
            if opt.isChecked():
                boneMembrane = opt.text()
                optionsChecked = True
        if optionsChecked:
            return "The site was grafted with "+boneMaterial+" and covered with a "+boneMembrane+" membrane. "
        else:
            return ""

    def generateParagraph(self):
        #print("generating paragraph")
        txt = ""
        txt += self.getIncision()
        txt += self.getFlap()
        txt += self.getExtraction()
        txt += self.getOsteotomy()

        txt += self.getBoneDensity()
        txt += self.getPerforation()
        txt += self.getTapped()
        txt += self.getSinus()
        txt += self.getSeated()
        txt += self.getLot()
        txt += self.getExpiration()
        txt += self.getDepth()
        #txt += self.getCover()
        txt += self.getHealingCap()
        txt += self.getGraft()
        return txt

    def getRestorativeParts(self):
        #print(self.restorativePartsList.findChildren(QCheckBox))
        options = self.findChild(QScrollArea, "restorativePartsScroll").findChildren(QCheckBox)

        txt = ""
        for opt in options:
            if opt.isChecked():
                if opt.text() not in txt:
                    txt += opt.text() + "\n"
        return txt

    def getImplant(self):
        if len(str(self.implant)) < 3:
            return ""

        return (str(self.implant) + " #"+ str(self.toothNumber.text()))


    def getExpiration(self):
        expDate = self.expirationDate.text()
        if expDate == QDate.currentDate().toString('M/d/yyyy'):
            return ""
        else:
            return "Expires: " + self.expirationDate.text() + ". "

    def getLot(self):
        if len(self.lotNumber.text()) < 2:
            return ""
        else:
            return "Lot #" + self.lotNumber.text() + ". "

    def getHealingCapList(self):
        if not self.healingCapBox.isChecked():
            return ""

        btn1 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton1')
        btn2 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton2')
        edit = self.healingCapBox.findChild(QLineEdit, 'healingCapEdit')


        if btn1.isChecked():
            return ""
        elif btn2.isChecked():
            return edit.text() + " healing cap #" + str(self.toothNumber.text())
        else:
            return ""

    def getImplantDetails(self):
        if len(str(self.implant)) < 3:
            return ""

        return (str(self.implant) + " #"+ str(self.toothNumber.text())) + "; "

    def getHealingCapDetails(self):
        if not self.healingCapBox.isChecked():
            return ""

        btn1 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton1')
        btn2 = self.healingCapBox.findChild(QRadioButton, 'healingCapButton2')
        edit = self.healingCapBox.findChild(QLineEdit, 'healingCapEdit')


        if btn1.isChecked():
            return ""
        elif btn2.isChecked():
            return edit.text() + " healing cap; "
        else:
            return ""

    def getExtractionDetails(self):
        if not self.extractionBox.isChecked():
            return ""

        edit = self.extractionBox.findChild(QLineEdit, "extractionEdit")
        btn1 = self.extractionBox.findChild(QRadioButton, "extractionButton1")
        btn2 = self.extractionBox.findChild(QRadioButton, "extractionButton2")
        btn3 = self.extractionBox.findChild(QRadioButton, "extractionButton3")
        btn4 = self.extractionBox.findChild(QRadioButton, "extractionButton4")

        txt = ""

        if btn1.isChecked():
            txt += "The socket was perfectly preserved; "
        elif btn2.isChecked():
            txt += "The socket was well preserved; "
        elif btn3.isChecked():
            txt += "The socket was moderately preserved; "
        elif btn4.isChecked():
            txt += "There was damage to the socket; "
        return txt

    def getGraftDetails(self):
        if not self.graftBox.isChecked():
            return ""

        boneMaterialBox = self.findChild(QGroupBox, "boneMaterialBox")
        boneMembraneBox = self.findChild(QGroupBox, "boneMembraneBox")
        boneMaterial = ""
        boneMembrane = ""

        optionsChecked = False
        for opt in boneMaterialBox.findChildren(QRadioButton):
            if opt.isChecked():
                boneMaterial = opt.text()
                optionsChecked = True
        for opt in boneMembraneBox.findChildren(QRadioButton):
            if opt.isChecked():
                boneMembrane = opt.text()
                optionsChecked = True
        if optionsChecked:
            return boneMaterial+" and "+boneMembrane+"; "
        else:
            return ""

    def getDetails(self):
        implant = self.getImplantDetails()
        cap = self.getHealingCapDetails()
        extraction = self.getExtractionDetails()
        graft = self.getGraftDetails()

        txt = implant + cap + graft + extraction

        return txt


