import application
from docx import Document
from docx.shared import Inches

doc = Document('application\\data\\template.docx')

tables = doc.tables
print(doc.get_merge_fields())

p = tables[0].rows[0].cells[0].add_paragraph()
r = p.add_run()
r.add_picture('C:\\PythonProjects\\ImplantApp\\app\\static\\images\\exampleXray.jpeg',width=Inches(2.5))

doc.save('test2.docx')
